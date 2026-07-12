#!/usr/bin/env python3
"""Casegraph — per-matter case file index and verification gates.

The legal-case-file analog of a code index: a persistent, deterministic index of
every document in a matter directory, plus machine-enforced gates that check
agent outputs against that index (citation resolution, cross-matter isolation,
staleness).

Design contract (see SPEC.md):
- The index lives INSIDE the matter directory (``<matter_dir>/.casegraph/``),
  outside this repo, so isolation is physical.
- Deterministic and provenance-first: hashes, structured headers, filename
  patterns. No inference. Unreadable content is flagged, never guessed.
- Contamination checks never read another matter's directory; cross-matter
  detection uses a salted-hash fingerprint store only.
- Gate commands exit non-zero on failure so skills/CI can chain them.

Stdlib-only core; pypdf / python-docx are optional (graceful degradation).
Synthetic data only in this repository. Attorney review required before any
real-matter use.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
import time
import unicodedata
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

SCHEMA_VERSION = 2  # v2: documents.jsonl rows carry declared_ranges
TOOL_VERSION = "1.2.0"
INDEX_DIRNAME = ".casegraph"
TEXT_CACHE_DIRNAME = "text"
OCR_QUEUE_FILENAME = "needs_ocr.json"
# text_extractable values that mean "no reliable text for cite/quote gates"
_OCR_NEEDED_STATUSES = frozenset({"none", "partial", "unsupported"})

# File types the indexer attempts text extraction for.
_TEXT_EXTS = {".txt", ".md", ".csv", ".log", ".json", ".yaml", ".yml", ".html", ".htm"}
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}
_EML_EXTS = {".eml"}

# Bates identifier parsing comes in two strictness levels:
#
# _BATES_TOKEN_RE (permissive) — filename/header parsing, where a space or
# underscore may separate prefix and number ("TVRR-PROD 000123", scan_...).
_BATES_TOKEN_RE = re.compile(
    r"\b([A-Z][A-Z0-9]{1,11}(?:-[A-Z][A-Z0-9]{1,11})*)[-_ ]?(\d{3,8})\b"
)

# _BATES_TEXT_RE (strict) — citation/isolation scanning of PROSE, where the
# permissive form drowns in false positives: "November 2024" reads as
# NOVEMBER-002024, "Section 218" as SECTION-000218, etc. Prose bates must be
# hyphen/underscore-joined (no bare space) AND look like a real production
# number: 5+ digits, or shorter with a leading zero (TVRR-PROD-000001 keeps
# matching; issue codes like DAM-001 are excluded separately via allowlisted
# prefixes, and 4-digit years never match).
_BATES_TEXT_RE = re.compile(
    r"\b([A-Z][A-Z0-9]{1,11}(?:-[A-Z][A-Z0-9]{1,11})*)[-_](0\d{2,7}|\d{5,8})\b"
)

# Structured header fields used by production documents / fixtures:
#   **Bates Range:** TVRR-PROD-000001 - TVRR-PROD-000004
_HEADER_FIELD_RE = re.compile(
    r"^\*\*(?P<key>[A-Za-z /-]+):\*\*\s*(?P<value>.+?)\s*$", re.MULTILINE
)

# Candidate person/org names in outputs: 2-4 capitalized words in sequence
# (allowing initials like "J.T." and connectors). High recall, moderate
# precision — used only for WARN-level findings, never FAIL.
# Connectors use [ \t]+ (NOT \s+): a candidate must never span a line break,
# or heading + first-prose-word fuse into junk like
# 'Bates Range Normalization\n\nBates' (receipt-run finding).
_NAME_CANDIDATE_RE = re.compile(
    r"\b([A-Z][a-zA-Z.]{1,20}(?:[ \t]+(?:of|the|and|for|de|van|von)[ \t]+|[ \t]+)"
    r"[A-Z][a-zA-Z.]{1,20}(?:[ \t]+[A-Z][a-zA-Z.]{1,20}){0,2})\b"
)

# Straight and curly double quotes (handoff quote checks must not miss
# typographic quotes common in Word/PDF extractions).
_QUOTE_SPLIT_RE = re.compile(r'["\u201c\u201d]')


def _iter_quoted_spans(text: str):
    """Yield the text inside double-quote pairs, line by line.

    Parity-based pairing: splitting a line on quote characters puts quoted
    content at ODD indices. A naive pair-regex with a minimum length skips a
    short quotation and then cross-pairs its CLOSING mark with the next
    quotation's OPENING mark, "verifying" the prose between two quotes
    (receipt-run finding). Quotes never span lines.
    """
    for line in text.splitlines():
        parts = _QUOTE_SPLIT_RE.split(line)
        for i in range(1, len(parts), 2):
            span = parts[i]
            if 1 <= len(span) <= 300:
                yield span


# ── small utilities ─────────────────────────────────────────────────────────

def _utcnow() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _sha256_file(path: Path, chunk: int = 1 << 20) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while True:
            b = f.read(chunk)
            if not b:
                break
            h.update(b)
    return h.hexdigest()


def _normalize_identifier(s: str) -> str:
    """Normalize a name/identifier for comparison and fingerprinting.

    NFKC (homoglyph/width defense) -> casefold -> strip punctuation ->
    collapse whitespace. ``J.T.`` and ``J T`` compare equal.
    """
    s = unicodedata.normalize("NFKC", s)
    s = s.casefold()
    s = re.sub(r"[^\w\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _salted_hash(value: str, salt: str) -> str:
    return hashlib.sha256((salt + "\x1f" + _normalize_identifier(value)).encode("utf-8")).hexdigest()


def _read_text_best_effort(path: Path) -> str:
    data = path.read_bytes()
    for enc in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return data.decode("utf-8", errors="replace")


def _index_dir(matter_dir: Path) -> Path:
    return matter_dir / INDEX_DIRNAME


def _ocr_needed_rows(rows: List[dict]) -> List[dict]:
    """Documents that need OCR / better text before content citation is safe."""
    out = []
    for r in rows:
        status = r.get("text_extractable") or "unsupported"
        if status not in _OCR_NEEDED_STATUSES:
            continue
        out.append({
            "relpath": r["relpath"],
            "sha256": r.get("sha256"),
            "ext": r.get("ext"),
            "text_extractable": status,
            "bates_prefix": r.get("bates_prefix"),
            "bates_start": r.get("bates_start"),
            "bates_end": r.get("bates_end"),
            "size": r.get("size"),
        })
    out.sort(key=lambda x: x["relpath"])
    return out


def write_ocr_queue(matter_dir: Path, rows: List[dict], matter_id: str) -> Path:
    """Persist OCR queue for agents: build continues, queue drives offline OCR."""
    needed = _ocr_needed_rows(rows)
    payload = {
        "schema_version": 1,
        "tool_version": TOOL_VERSION,
        "matter_id": matter_id,
        "generated_at": _utcnow(),
        "count": len(needed),
        "documents": needed,
        "guidance": (
            "OCR these into 01_production/text/ (or add a searchable text layer "
            "to the PDF via ocrmypdf), then re-run: casegraph build. Prefer "
            "background terminal for large sets. Do not send scans to a remote "
            "vision API without PROVIDER_AUTH."
        ),
    }
    path = _index_dir(matter_dir) / OCR_QUEUE_FILENAME
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(payload, indent=2, sort_keys=True), encoding="utf-8")
    os.replace(tmp, path)
    return path


def cmd_export_ocr_queue(args) -> int:
    """Emit the OCR queue (from last build artifact, or recompute from index)."""
    matter_dir = Path(args.matter_dir).resolve()
    manifest = load_manifest(matter_dir)
    rows = load_documents(matter_dir)
    queue_path = write_ocr_queue(matter_dir, rows, manifest["matter_id"])
    needed = _ocr_needed_rows(rows)
    if args.json:
        print(json.dumps({
            "matter_id": manifest["matter_id"],
            "path": str(queue_path),
            "count": len(needed),
            "documents": needed,
        }, indent=2))
    else:
        print(f"OCR queue: {len(needed)} document(s) need text/OCR "
              f"(written to {queue_path})")
        for d in needed[:50]:
            print(f"  - [{d['text_extractable']}] {d['relpath']}")
        if len(needed) > 50:
            print(f"  … and {len(needed) - 50} more")
        if needed:
            print("After OCR into 01_production/text/ (or searchable PDF), re-run: "
                  "casegraph build")
    return 0 if not needed else 1


def _manifest_path(matter_dir: Path) -> Path:
    return _index_dir(matter_dir) / "manifest.json"


def _documents_path(matter_dir: Path) -> Path:
    return _index_dir(matter_dir) / "documents.jsonl"


def _entities_path(matter_dir: Path) -> Path:
    return _index_dir(matter_dir) / "entities.json"


def _chronology_path(matter_dir: Path) -> Path:
    return _index_dir(matter_dir) / "chronology.jsonl"


def load_manifest(matter_dir: Path) -> dict:
    p = _manifest_path(matter_dir)
    if not p.exists():
        raise SystemExit(
            f"ERROR: no casegraph index at {p}. Run: casegraph.py init {matter_dir} "
            f"--matter-id <ID> --bates-prefix <PREFIX>"
        )
    return json.loads(p.read_text(encoding="utf-8"))


def save_manifest(matter_dir: Path, manifest: dict) -> None:
    manifest["updated"] = _utcnow()
    p = _manifest_path(matter_dir)
    p.parent.mkdir(parents=True, exist_ok=True)
    tmp = p.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(manifest, indent=2, sort_keys=True), encoding="utf-8")
    os.replace(tmp, p)


def load_documents(matter_dir: Path) -> List[dict]:
    p = _documents_path(matter_dir)
    if not p.exists():
        return []
    rows = []
    for line in p.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line:
            rows.append(json.loads(line))
    return rows


def save_documents(matter_dir: Path, rows: List[dict]) -> None:
    p = _documents_path(matter_dir)
    p.parent.mkdir(parents=True, exist_ok=True)
    tmp = p.with_suffix(".jsonl.tmp")
    with open(tmp, "w", encoding="utf-8", newline="\n") as f:
        for row in sorted(rows, key=lambda r: r["relpath"]):
            f.write(json.dumps(row, sort_keys=True) + "\n")
    os.replace(tmp, p)


def load_entities(matter_dir: Path) -> dict:
    p = _entities_path(matter_dir)
    if not p.exists():
        return {}
    return json.loads(p.read_text(encoding="utf-8"))


def save_entities(matter_dir: Path, entities: dict) -> None:
    p = _entities_path(matter_dir)
    p.parent.mkdir(parents=True, exist_ok=True)
    tmp = p.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(entities, indent=2, sort_keys=True), encoding="utf-8")
    os.replace(tmp, p)


# ── text extraction (graceful degradation, no inference) ───────────────────

def extract_text(path: Path) -> Tuple[Optional[str], str, Optional[int]]:
    """Return (text, extractable_status, pages).

    extractable_status: "full" | "partial" | "none" | "unsupported".
    Never guesses: a PDF page with no text layer contributes nothing and
    downgrades status to partial/none.
    """
    ext = path.suffix.lower()
    if ext in _TEXT_EXTS:
        return _read_text_best_effort(path), "full", None
    if ext in _PDF_EXTS:
        try:
            from pypdf import PdfReader
        except ImportError:
            return None, "unsupported", None
        try:
            reader = PdfReader(str(path))
            pages = len(reader.pages)
            texts, empty = [], 0
            for pg in reader.pages:
                t = (pg.extract_text() or "").strip()
                if t:
                    texts.append(t)
                else:
                    empty += 1
            if not texts:
                return None, "none", pages
            status = "partial" if empty else "full"
            return "\n\n".join(texts), status, pages
        except Exception:
            return None, "none", None
    if ext in _DOCX_EXTS:
        try:
            import docx  # python-docx
        except ImportError:
            return None, "unsupported", None
        try:
            d = docx.Document(str(path))
            parts = [p.text for p in d.paragraphs]
            for table in d.tables:
                for row in table.rows:
                    parts.append("\t".join(c.text for c in row.cells))
            text = "\n".join(p for p in parts if p)
            return (text, "full", None) if text.strip() else (None, "none", None)
        except Exception:
            return None, "none", None
    if ext in _EML_EXTS:
        try:
            import email
            from email import policy
            msg = email.message_from_bytes(path.read_bytes(), policy=policy.default)
            headers = "\n".join(
                f"{k}: {msg.get(k, '')}" for k in ("From", "To", "Cc", "Date", "Subject")
            )
            body = msg.get_body(preferencelist=("plain", "html"))
            body_text = body.get_content() if body else ""
            return headers + "\n\n" + str(body_text), "full", None
        except Exception:
            return None, "none", None
    return None, "unsupported", None


def parse_header_fields(text: str) -> Dict[str, str]:
    """Parse the ``**Field:** value`` structured header convention."""
    fields = {}
    for m in _HEADER_FIELD_RE.finditer(text[:4000]):
        key = m.group("key").strip().lower().replace(" ", "_").replace("/", "_")
        fields[key] = m.group("value").strip()
    return fields


def parse_bates_range(value: str) -> Optional[Tuple[str, int, int]]:
    """Parse 'TVRR-PROD-000001 - TVRR-PROD-000004' (also 'through', en-dash).

    Returns (prefix, start, end) or None.
    """
    tokens = _BATES_TOKEN_RE.findall(value.upper())
    if not tokens:
        return None
    prefix = tokens[0][0]
    nums = [int(n) for p, n in tokens if p == prefix]
    if not nums:
        return None
    return prefix, min(nums), max(nums)


# Declared ranges in a production cover letter: "TVRR-PROD-000005 through
# 000016", "TVRR-PROD-000061 to TVRR-PROD-000071", en/em-dash forms.
_DECLARED_RANGE_RE = re.compile(
    r"\b([A-Z][A-Z0-9]{1,11}(?:-[A-Z][A-Z0-9]{1,11})*)-(0\d{2,7}|\d{5,8})"
    r"\s*(?:through|to|thru|[–—-])\s*"
    r"(?:[A-Z][A-Z0-9-]{1,23}-)?(0\d{2,7}|\d{5,8})\b",
    re.IGNORECASE,
)


def parse_declared_ranges(text: str) -> List[Tuple[str, int, int]]:
    """Bates ranges DECLARED by a production cover letter / index.

    Declared ranges support gap analysis (inventory vs files on disk). They do
    NOT ground fact citations by default — citing a Bates that is only
    declared, not indexed, is a FAIL unless ``--allow-declared-gaps`` is set
    (intentional gap-analysis packages). Numbers OUTSIDE every declared range
    remain hard failures either way.
    """
    out: List[Tuple[str, int, int]] = []
    for m in _DECLARED_RANGE_RE.finditer(text.upper()):
        start, end = int(m.group(2)), int(m.group(3))
        if end >= start:
            out.append((m.group(1), start, end))
    return out


def bates_from_filename(name: str) -> Optional[Tuple[str, int]]:
    # Underscores are word characters, so "scan_TVRR-PROD-000010.pdf" would
    # otherwise never get a word boundary before TVRR and mis-parse as
    # prefix "PROD". Treat underscores as separators for filename parsing.
    m = _BATES_TOKEN_RE.search(name.upper().replace("_", " "))
    if m:
        return m.group(1), int(m.group(2))
    return None


# ── init / build / status ───────────────────────────────────────────────────

def cmd_init(args) -> int:
    matter_dir = Path(args.matter_dir).resolve()
    matter_dir.mkdir(parents=True, exist_ok=True)
    mp = _manifest_path(matter_dir)
    if mp.exists() and not args.force:
        print(f"ERROR: index already exists at {mp} (use --force to reinitialize)")
        return 2
    manifest = {
        "schema_version": SCHEMA_VERSION,
        "tool_version": TOOL_VERSION,
        "matter_id": args.matter_id,
        "bates_prefixes": sorted({p.upper() for p in (args.bates_prefix or [])}),
        "created": _utcnow(),
        "counts": {"documents": 0},
    }
    save_manifest(matter_dir, manifest)
    print(f"Initialized casegraph for matter '{args.matter_id}' at {_index_dir(matter_dir)}")
    if not manifest["bates_prefixes"]:
        print("NOTE: no --bates-prefix registered; isolation checks on bates will flag everything.")
    return 0


def _iter_matter_files(matter_dir: Path) -> Iterable[Path]:
    for root, dirs, files in os.walk(matter_dir):
        dirs[:] = [d for d in dirs if d != INDEX_DIRNAME and not d.startswith(".")]
        for name in files:
            if name.startswith("."):
                continue
            yield Path(root) / name


def _scan_file(matter_dir: Path, path: Path, no_text_cache: bool) -> dict:
    rel = path.relative_to(matter_dir).as_posix()
    st = path.stat()
    sha = _sha256_file(path)
    row = {
        "relpath": rel,
        "sha256": sha,
        "size": st.st_size,
        "mtime_iso": datetime.fromtimestamp(st.st_mtime, tz=timezone.utc).strftime(
            "%Y-%m-%dT%H:%M:%SZ"
        ),
        "ext": path.suffix.lower(),
        "indexed_at": _utcnow(),
        "pages": None,
        "text_extractable": "unsupported",
        "bates_prefix": None,
        "bates_start": None,
        "bates_end": None,
        "doc_date": None,
        "author": None,
        "custodian": None,
        "doc_type": None,
        "title": None,
        "dupes_of": None,
        "declared_ranges": [],
    }
    text, status, pages = extract_text(path)
    row["text_extractable"] = status
    row["pages"] = pages
    if text:
        fields = parse_header_fields(text)
        if "bates_range" in fields:
            parsed = parse_bates_range(fields["bates_range"])
            if parsed:
                row["bates_prefix"], row["bates_start"], row["bates_end"] = parsed
        for src, dst in (
            ("date", "doc_date"), ("author", "author"), ("custodian", "custodian"),
            ("document_type", "doc_type"), ("document_id", "title"),
        ):
            if fields.get(src):
                row[dst] = fields[src]
        # Only explicit Document Type headers may declare production ranges.
        # Filename / title alone ("cover_letter.md") must NOT harvest ranges —
        # that was a cite-laundering vector (red-team P1-6).
        _kind = (row.get("doc_type") or "").strip().lower()
        if "cover letter" in _kind or "production index" in _kind:
            row["declared_ranges"] = [list(r) for r in parse_declared_ranges(text)]
        if not no_text_cache:
            cache = _index_dir(matter_dir) / TEXT_CACHE_DIRNAME
            cache.mkdir(parents=True, exist_ok=True)
            (cache / f"{sha}.txt").write_text(text, encoding="utf-8", newline="\n")
    if row["bates_start"] is None:
        fb = bates_from_filename(path.name)
        if fb:
            row["bates_prefix"], row["bates_start"] = fb
            row["bates_end"] = fb[1]
    return row


def _harvest_header_entities(matter_dir: Path, rows: List[dict]) -> int:
    """Register Author/Custodian header values as entities (origin=header)."""
    entities = load_entities(matter_dir)
    added = 0
    for row in rows:
        for field, role in (("author", "author"), ("custodian", "custodian")):
            raw = row.get(field)
            if not raw:
                continue
            # Header values like "R.K., Trainmaster, Test Valley Railroad" carry
            # name + role + org; register each comma part as its own entity.
            for part in [p.strip() for p in raw.split(",") if p.strip()]:
                key = _normalize_identifier(part)
                if not key or len(key) < 2:
                    continue
                ent = entities.setdefault(
                    key, {"display": part, "aliases": [], "role": role,
                          "origin": "header", "sources": {}}
                )
                ent["sources"][row["relpath"]] = ent["sources"].get(row["relpath"], 0) + 1
                added += 1
    save_entities(matter_dir, entities)
    return added


def cmd_build(args) -> int:
    matter_dir = Path(args.matter_dir).resolve()
    manifest = load_manifest(matter_dir)
    old_rows = {r["relpath"]: r for r in load_documents(matter_dir)}
    # Index built by an older tool schema: unchanged files would keep stale
    # rows missing newer fields (e.g. declared_ranges), so force a full
    # re-scan once and stamp the new schema.
    if manifest.get("schema_version") != SCHEMA_VERSION:
        old_rows = {}
        manifest["schema_version"] = SCHEMA_VERSION
    new_rows: List[dict] = []
    n_new = n_changed = n_same = 0

    for path in _iter_matter_files(matter_dir):
        rel = path.relative_to(matter_dir).as_posix()
        st = path.stat()
        old = old_rows.get(rel)
        if (
            old is not None
            and old["size"] == st.st_size
            and old["mtime_iso"] == datetime.fromtimestamp(
                st.st_mtime, tz=timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        ):
            new_rows.append(old)
            n_same += 1
            continue
        row = _scan_file(matter_dir, path, args.no_text_cache)
        if old is None:
            n_new += 1
        else:
            n_changed += 1
        new_rows.append(row)

    removed = sorted(set(old_rows) - {r["relpath"] for r in new_rows})

    # Exact-duplicate marking: first relpath (sorted) is canonical.
    by_sha: Dict[str, List[dict]] = {}
    for r in new_rows:
        by_sha.setdefault(r["sha256"], []).append(r)
    n_dupes = 0
    for sha, group in by_sha.items():
        group.sort(key=lambda r: r["relpath"])
        for extra in group[1:]:
            extra["dupes_of"] = group[0]["relpath"]
            n_dupes += 1
        if group[0].get("dupes_of"):
            group[0]["dupes_of"] = None

    save_documents(matter_dir, new_rows)
    n_entities = _harvest_header_entities(matter_dir, new_rows)

    # Bates coverage report (registered prefixes only).
    coverage_notes: List[str] = []
    for prefix in manifest.get("bates_prefixes", []):
        nums: List[Tuple[int, int]] = sorted(
            (r["bates_start"], r["bates_end"])
            for r in new_rows
            if r.get("bates_prefix") == prefix and r.get("bates_start") is not None
        )
        prev_end = None
        for start, end in nums:
            if prev_end is not None and start <= prev_end:
                coverage_notes.append(f"{prefix}: overlap at {start:06d} (prev range ends {prev_end:06d})")
            if prev_end is not None and start > prev_end + 1:
                coverage_notes.append(f"{prefix}: gap {prev_end + 1:06d}-{start - 1:06d}")
            prev_end = max(prev_end or 0, end)

    unreadable = [r["relpath"] for r in new_rows if r["text_extractable"] in ("none",)]
    ocr_queue_path = write_ocr_queue(matter_dir, new_rows, manifest["matter_id"])
    ocr_needed = _ocr_needed_rows(new_rows)
    manifest["counts"] = {
        "documents": len(new_rows),
        "duplicates": n_dupes,
        "unreadable": len(unreadable),
        "ocr_needed": len(ocr_needed),
        "entities": len(load_entities(matter_dir)),
    }
    save_manifest(matter_dir, manifest)

    report = {
        "matter_id": manifest["matter_id"],
        "documents": len(new_rows),
        "new": n_new, "changed": n_changed, "unchanged": n_same,
        "removed": removed,
        "duplicates": n_dupes,
        "unreadable": unreadable,
        "ocr_needed": len(ocr_needed),
        "ocr_queue": str(ocr_queue_path),
        "bates_coverage_notes": coverage_notes,
        "entity_mentions_registered": n_entities,
    }
    if args.json:
        print(json.dumps(report, indent=2))
    else:
        print(f"Indexed {len(new_rows)} documents for matter '{manifest['matter_id']}' "
              f"({n_new} new, {n_changed} changed, {n_same} unchanged"
              f"{', ' + str(len(removed)) + ' removed' if removed else ''}).")
        if n_dupes:
            print(f"Exact duplicates: {n_dupes} (marked dupes_of canonical copy)")
        if unreadable:
            print(f"UNREADABLE (no text layer — manual/OCR review needed): {len(unreadable)}")
            for u in unreadable[:20]:
                print(f"  - {u}")
        if ocr_needed:
            print(f"OCR QUEUE: {len(ocr_needed)} doc(s) → {ocr_queue_path}")
            print("  Export anytime: casegraph export-ocr-queue <matter_dir>")
            print("  After OCR into 01_production/text/ (or searchable PDF), re-run: "
                  "casegraph build")
        for note in coverage_notes:
            print(f"BATES: {note}")
    return 0


def cmd_status(args) -> int:
    matter_dir = Path(args.matter_dir).resolve()
    manifest = load_manifest(matter_dir)
    old_rows = {r["relpath"]: r for r in load_documents(matter_dir)}
    added, changed = [], []
    seen = set()
    for path in _iter_matter_files(matter_dir):
        rel = path.relative_to(matter_dir).as_posix()
        seen.add(rel)
        old = old_rows.get(rel)
        if old is None:
            added.append(rel)
            continue
        st = path.stat()
        mt = datetime.fromtimestamp(st.st_mtime, tz=timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        if old["size"] != st.st_size or old["mtime_iso"] != mt:
            changed.append(rel)
        elif args.deep and _sha256_file(path) != old["sha256"]:
            changed.append(rel)
    removed = sorted(set(old_rows) - seen)
    stale = bool(added or changed or removed)
    report = {
        "matter_id": manifest["matter_id"],
        "stale": stale,
        "added": sorted(added), "changed": sorted(changed), "removed": removed,
        "indexed_documents": len(old_rows),
    }
    if args.json:
        print(json.dumps(report, indent=2))
    else:
        if stale:
            print(f"STALE index for matter '{manifest['matter_id']}':")
            for label, items in (("added", added), ("changed", changed), ("removed", removed)):
                for it in sorted(items):
                    print(f"  {label}: {it}")
            print("Run `casegraph.py build` before relying on the index.")
        else:
            print(f"Index current: {len(old_rows)} documents, matter '{manifest['matter_id']}'.")
    return 1 if stale else 0


# ── query ────────────────────────────────────────────────────────────────────

def _resolve_bates(rows: List[dict], prefix: str, number: int) -> Optional[dict]:
    for r in rows:
        if (
            r.get("bates_prefix") == prefix
            and r.get("bates_start") is not None
            and r["bates_start"] <= number <= (r.get("bates_end") or r["bates_start"])
        ):
            return r
    return None


def cmd_query(args) -> int:
    matter_dir = Path(args.matter_dir).resolve()
    load_manifest(matter_dir)
    rows = load_documents(matter_dir)
    results: List[dict] = []

    if args.bates:
        m = _BATES_TOKEN_RE.search(args.bates.upper())
        if not m:
            print(f"ERROR: not a bates identifier: {args.bates}")
            return 2
        hit = _resolve_bates(rows, m.group(1), int(m.group(2)))
        results = [hit] if hit else []
    elif args.doc:
        needle = args.doc.replace("\\", "/").lower()
        results = [r for r in rows if needle in r["relpath"].lower()]
    elif args.entity:
        entities = load_entities(matter_dir)
        key = _normalize_identifier(args.entity)
        ent = entities.get(key)
        if ent is None:
            for k, v in entities.items():
                if key in k or any(key == _normalize_identifier(a) for a in v.get("aliases", [])):
                    ent = v
                    break
        print(json.dumps(ent or {}, indent=2))
        return 0 if ent else 1
    elif args.grep:
        pattern = re.compile(args.grep, re.IGNORECASE)
        cache = _index_dir(matter_dir) / TEXT_CACHE_DIRNAME
        for r in rows:
            fp = cache / f"{r['sha256']}.txt"
            if not fp.exists():
                continue
            for i, line in enumerate(fp.read_text(encoding="utf-8").splitlines(), 1):
                if pattern.search(line):
                    results.append({"relpath": r["relpath"], "line": i, "text": line.strip()[:240]})
    else:
        results = rows

    print(json.dumps(results, indent=2))
    return 0 if results else 1


# ── verification gates ──────────────────────────────────────────────────────

def _allowlist_path() -> Path:
    return Path(__file__).resolve().parent.parent / "data" / "legal_allowlist.txt"


def _load_allowlist() -> set:
    """Global legal allowlist (courts, statutes, common legal phrases) shipped
    with the skill. Normalized entries. Missing file is a hard error for
    isolation semantics — an empty allowlist silently WARNs on every legal
    phrase and diverges from SPEC."""
    path = _allowlist_path()
    if not path.exists():
        raise FileNotFoundError(
            f"casegraph allowlist missing: {path} — restore "
            f"skills/legal/casegraph/data/legal_allowlist.txt"
        )
    entries: set = set()
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line and not line.startswith("#"):
            entries.add(_normalize_identifier(line))
    if not entries:
        raise ValueError(f"casegraph allowlist is empty: {path}")
    return entries


def _extract_citations(text: str) -> List[Tuple[str, int]]:
    """Bates citations in PROSE — uses the strict text-form regex.

    The permissive filename form would turn 'November 2024' into
    NOVEMBER-002024 and 'Section 218' into SECTION-000218 (receipt-run
    finding); prose citations require the joined, production-number shape.
    """
    out = []
    for m in _BATES_TEXT_RE.finditer(text.upper()):
        out.append((m.group(1), int(m.group(2))))
    return out


# Lines carrying these markers discuss documents that are EXPECTED TO BE
# ABSENT (production-gap analysis). An unresolved citation there is the
# point of the sentence, not a fabrication.
_GAP_MARKER_RE = re.compile(
    r"not (?:been )?produced|missing|gap|not searched|no documents produced|"
    r"outstanding|withheld|expected but|claimed not to exist", re.IGNORECASE
)

# Quoted skill/gate meta-language is template text, not a document quotation.
# Match at START only — appending "requires attorney review" to a fabricated
# quote must not exempt it (red-team P1-5).
_META_QUOTE_RE = re.compile(
    r"^(?:evidence (?:suggests|supports|contradicts)|requires attorney review|"
    r"attorney review required)\b", re.IGNORECASE
)


def cmd_verify_cites(args) -> int:
    matter_dir = Path(args.matter_dir).resolve()
    manifest = load_manifest(matter_dir)
    rows = load_documents(matter_dir)
    output_path = Path(args.output_file)
    if not output_path.exists():
        print(f"ERROR: output file not found: {output_path}")
        return 2
    text = _read_text_best_effort(output_path)

    registered = set(manifest.get("bates_prefixes", []))
    declared: List[Tuple[str, int, int]] = [
        (r[0], int(r[1]), int(r[2]))
        for doc in rows for r in (doc.get("declared_ranges") or [])
    ]

    def _in_declared(prefix: str, number: int) -> bool:
        return any(p == prefix and s <= number <= e for p, s, e in declared)

    failures: List[str] = []
    gap_notes: List[str] = []
    checked = 0
    seen: set = set()
    # Line-aware pass so gap-analysis context is visible per citation:
    # a citation on a "not produced / missing / gap" line is EXPECTED to be
    # unresolved — that's the finding being reported, not a fabrication.
    for line in text.splitlines():
        gap_context = bool(_GAP_MARKER_RE.search(line))
        for prefix, number in _extract_citations(line):
            if prefix not in registered:
                continue  # foreign prefixes are the isolation gate's job
            key = (prefix, number)
            if key in seen:
                continue
            seen.add(key)
            checked += 1
            resolved = _resolve_bates(rows, prefix, number) is not None
            allow_declared = getattr(args, "allow_declared_gaps", False)
            if not resolved and _in_declared(prefix, number):
                msg = (
                    f"declared-not-indexed: {prefix}-{number:06d} is inside a "
                    f"range the production cover letter declares, but no file "
                    f"for it is indexed"
                )
                if allow_declared:
                    # Explicit gap-analysis opt-in: INFO, not a fabrication FAIL.
                    gap_notes.append(
                        f"{msg} — allowed via --allow-declared-gaps "
                        f"(document not in reviewed set)"
                    )
                else:
                    # Default fail-closed: cover-letter text must not launder
                    # fact cites for documents that were never indexed.
                    failures.append(
                        f"{msg} (use --allow-declared-gaps only for intentional "
                        f"gap-analysis packages)"
                    )
            elif not resolved and gap_context:
                gap_notes.append(
                    f"gap-context citation (expected absent): {prefix}-{number:06d}"
                )
            elif not resolved:
                failures.append(f"unresolved citation: {prefix}-{number:06d} "
                                f"(no indexed document covers this number)")
            elif resolved and gap_context:
                gap_notes.append(
                    f"NOTE: {prefix}-{number:06d} is listed as missing/not "
                    f"produced but RESOLVES in the index — verify the gap claim"
                )

    # Fail-closed: a review/handoff package with zero same-matter citations is
    # not a vacuous PASS. Opt out only with --allow-empty (intake drafts, etc.).
    if checked == 0 and not getattr(args, "allow_empty", False):
        failures.append(
            "no same-matter Bates citations found in output "
            "(vacuous verify-cites PASS refused; use --allow-empty only for "
            "non-review drafts that intentionally cite nothing)"
        )

    quote_misses: List[str] = []
    quotes_checked = 0
    # --quotes is the handoff default; --no-quotes opts out for draft passes.
    do_quotes = getattr(args, "quotes", True) and not getattr(args, "no_quotes", False)
    if do_quotes:
        cache = _index_dir(matter_dir) / TEXT_CACHE_DIRNAME
        # Collect pending quotes first, then scan one document at a time so
        # large matters never join the full corpus into a single blob.
        pending: Dict[str, str] = {}
        for q in _iter_quoted_spans(text):
            q_cmp = q.strip().strip(".… ").strip()
            # Short quotes still matter ("No handbrakes.") — floor is 12,
            # not 20 (red-team P1-4).
            if len(q_cmp) < 12:
                continue
            # Meta gate language only when it LEADS a short template phrase —
            # appending "requires attorney review" must not exempt a fabrication.
            if _META_QUOTE_RE.match(q_cmp) and len(q_cmp) < 60:
                continue
            quotes_checked += 1
            pending[_normalize_identifier(q_cmp)] = q[:120]
        for r in rows:
            if not pending:
                break
            fp = cache / f"{r['sha256']}.txt"
            if not fp.exists():
                continue
            doc = _normalize_identifier(fp.read_text(encoding="utf-8"))
            for qn in list(pending):
                if qn in doc:
                    del pending[qn]
        quote_misses = list(pending.values())

    report = {
        "output_file": str(output_path),
        "matter_id": manifest["matter_id"],
        "citations_checked": checked,
        "citation_failures": failures,
        "gap_notes": gap_notes,
        "quotes_checked": quotes_checked,
        "quote_misses": quote_misses,
        "pass": not failures and not quote_misses,
    }
    if args.json:
        print(json.dumps(report, indent=2))
    else:
        print(f"verify-cites: {checked} citations checked against matter "
              f"'{manifest['matter_id']}'.")
        for f_ in failures:
            print(f"  FAIL: {f_}")
        for note in gap_notes:
            print(f"  INFO: {note}")
        if do_quotes:
            print(f"verify-cites: {quotes_checked} quotes (>=12 chars) checked.")
            for q in quote_misses:
                print(f"  FAIL: quote not found in any indexed document: \"{q}...\"")
        if report["pass"]:
            print("PASS")
        else:
            print("FAIL")
    return 0 if report["pass"] else 1


def _fingerprint_store_load(path: Path) -> dict:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def cmd_export_fingerprint(args) -> int:
    matter_dir = Path(args.matter_dir).resolve()
    manifest = load_manifest(matter_dir)
    entities = load_entities(matter_dir)
    store_path = Path(args.store)
    store = _fingerprint_store_load(store_path)
    existing = store.get(manifest["matter_id"], {})
    salt = existing.get("salt") or hashlib.sha256(os.urandom(32)).hexdigest()[:32]

    entity_hashes = sorted({
        _salted_hash(k, salt) for k in entities
    } | {
        _salted_hash(a, salt)
        for v in entities.values() for a in v.get("aliases", [])
    })
    prefix_hashes = sorted(_salted_hash(p, salt) for p in manifest.get("bates_prefixes", []))
    store[manifest["matter_id"]] = {
        "salt": salt,
        "bates_prefix_hashes": prefix_hashes,
        "entity_hashes": entity_hashes,
        "exported": _utcnow(),
    }
    store_path.parent.mkdir(parents=True, exist_ok=True)
    tmp = store_path.with_suffix(store_path.suffix + ".tmp")
    tmp.write_text(json.dumps(store, indent=2, sort_keys=True), encoding="utf-8")
    os.replace(tmp, store_path)
    print(f"Exported fingerprint for matter '{manifest['matter_id']}' "
          f"({len(entity_hashes)} entity hashes, {len(prefix_hashes)} prefix hashes) "
          f"to {store_path}")
    return 0


def cmd_check_isolation(args) -> int:
    matter_dir = Path(args.matter_dir).resolve()
    manifest = load_manifest(matter_dir)
    output_path = Path(args.output_file)
    if not output_path.exists():
        print(f"ERROR: output file not found: {output_path}")
        return 2
    text = _read_text_best_effort(output_path)
    matter_id = manifest["matter_id"]
    registered_prefixes = set(manifest.get("bates_prefixes", []))
    allowlist = _load_allowlist()
    entities = load_entities(matter_dir)
    known_keys = set(entities)
    for v in entities.values():
        for a in v.get("aliases", []):
            known_keys.add(_normalize_identifier(a))

    failures: List[str] = []
    warnings: List[str] = []

    # 1) Foreign bates prefixes → FAIL (high-precision cross-matter signal).
    foreign_seen: set = set()
    for prefix, number in _extract_citations(text):
        if prefix in registered_prefixes or prefix in foreign_seen:
            continue
        norm_prefix = _normalize_identifier(prefix)
        if norm_prefix in allowlist:
            continue
        foreign_seen.add(prefix)
        failures.append(
            f"foreign bates prefix '{prefix}' (e.g. {prefix}-{number:06d}) — not "
            f"registered to matter '{matter_id}'"
        )

    # 2) Fingerprint store — identifiers registered to OTHER matters → FAIL.
    fp_hits: List[str] = []
    if args.fingerprints:
        store = _fingerprint_store_load(Path(args.fingerprints))
        candidates = {c for c in _candidate_names(text, subspans=True)}
        candidates |= {p for p, _ in _extract_citations(text)}
        for other_id, entry in store.items():
            if other_id == matter_id:
                continue
            salt = entry.get("salt", "")
            hashes = set(entry.get("entity_hashes", [])) | set(entry.get("bates_prefix_hashes", []))
            for cand in candidates:
                norm = _normalize_identifier(cand)
                if not norm or norm in allowlist or norm in known_keys:
                    continue
                if _salted_hash(cand, salt) in hashes:
                    fp_hits.append(f"'{cand}' matches an identifier registered to matter "
                                   f"'{other_id}'")
        failures.extend(sorted(set(fp_hits)))

    # 3) Unregistered candidate names → WARN (moderate precision; attorney list).
    unknown: List[str] = []
    for cand in sorted(set(_candidate_names(text))):
        norm = _normalize_identifier(cand)
        if not norm or norm in allowlist or norm in known_keys:
            continue
        if any(norm in k or k in norm for k in known_keys):
            continue  # partial alias overlap — treat as known
        unknown.append(cand)
    warnings.extend(f"unregistered name in output: '{u}'" for u in unknown)

    passed = not failures and (not args.strict or not warnings)
    report = {
        "output_file": str(output_path),
        "matter_id": matter_id,
        "failures": failures,
        "warnings": warnings,
        "strict": bool(args.strict),
        "pass": passed,
    }
    if args.json:
        print(json.dumps(report, indent=2))
    else:
        print(f"check-isolation: matter '{matter_id}', output {output_path.name}")
        for f_ in failures:
            print(f"  FAIL: {f_}")
        for w in warnings:
            print(f"  WARN: {w}")
        print("PASS" if passed else "FAIL" if failures else "FAIL (strict: unresolved WARNs)")
    return 0 if passed else 1


def _candidate_names(text: str, subspans: bool = False) -> Iterable[str]:
    """Yield candidate person/org names from output text.

    With ``subspans=True``, also yield every contiguous 2+-word sub-span of
    each candidate — a maximal span like 'Witness Marcus Ellery' must still
    match a fingerprint registered as 'Marcus Ellery'. Used for the
    fingerprint check (recall matters); WARN reporting uses maximal spans only
    (readability matters).
    """
    # NFKC first: fullwidth/compatibility homoglyphs (e.g. 'Ｍarcus') must
    # fold to ASCII BEFORE the [A-Z]-anchored candidate regex runs, or a
    # homoglyph-spelled name evades extraction entirely (red-team finding).
    text = unicodedata.normalize("NFKC", text)
    # Strip code fences and markdown emphasis to reduce false positives.
    text = re.sub(r"```.*?```", " ", text, flags=re.DOTALL)
    text = text.replace("**", " ")
    seen: set = set()
    for line in text.splitlines():
        stripped = line.lstrip()
        # Structural markdown is template scaffolding, not prose that could
        # smuggle another matter's names to a reader: headings and table rows
        # produced 200+ Title-Case WARNs per package (receipt-run finding).
        if stripped.startswith("#") or stripped.startswith("|"):
            continue
        # Field-label lines ("Document Type: …", "- Attorney Review Flag: …"):
        # scan only the value AFTER the label colon, so template field names
        # don't WARN but a name in the value still does.
        label = re.match(r"^\s*(?:[-*]\s*)?[A-Z][A-Za-z ./-]{0,40}:\s*(.*)$", line)
        if label:
            line = label.group(1)
        for m in _NAME_CANDIDATE_RE.finditer(line):
            cand = m.group(1).strip()
            # Skip single-token ALL-CAPS (heading noise). Multi-word ALL-CAPS
            # like "MARCUS ELLERY" must still be checked (red-team P1-5).
            if cand.isupper() and len(cand.split()) < 2:
                continue
            if cand not in seen:
                seen.add(cand)
                yield cand
            if subspans:
                words = cand.split()
                for width in range(2, len(words)):
                    for i in range(len(words) - width + 1):
                        sub = " ".join(words[i:i + width])
                        if sub not in seen:
                            seen.add(sub)
                            yield sub


_CHRONO_DATE_RE = re.compile(r"^\s*(?:[-*]\s*)?(?:\*\*)?Date(?:\*\*)?:\s*(?:\*\*)?"
                             r"\[?(\d{4}-\d{2}-\d{2})\]?", re.MULTILINE)

# Table-layout chronology rows: | 2024-11-12 | event… | TVRR-PROD-000001 |
# Without this, a package whose chronology is a markdown table verifies ZERO
# rows and PASSes vacuously (receipt-run finding).
_CHRONO_TABLE_ROW_RE = re.compile(
    r"^\s*\|\s*(?:\*\*)?(\d{4}-\d{2}-\d{2})(?:\*\*)?\s*\|(.+)$", re.MULTILINE)

_MONTHS = ["January", "February", "March", "April", "May", "June", "July",
           "August", "September", "October", "November", "December"]


def _date_variants(iso: str) -> List[str]:
    """Common renderings of an ISO date as they appear in source documents."""
    y, m, d = iso.split("-")
    mi, di = int(m), int(d)
    month = _MONTHS[mi - 1]
    return [
        iso,
        f"{month} {di}, {y}",
        f"{month} {di} {y}",
        f"{month[:3]} {di}, {y}",
        f"{m}/{d}/{y}", f"{mi}/{di}/{y}",
        f"{m}-{d}-{y}", f"{mi}-{di}-{y}",
        f"{m}/{d}/{y[2:]}", f"{mi}/{di}/{y[2:]}",
    ]


def cmd_verify_chronology(args) -> int:
    """Verify chronology entries: each dated event's Source citation must
    resolve, and the date must appear (in a common rendering) in the cited
    document's text.

    Semantics: unresolved citation -> FAIL; date absent from a readable cited
    doc -> WARN (dates are sometimes legitimately inferred — attorney list;
    --strict escalates); cited doc unreadable -> WARN (cannot verify).
    """
    matter_dir = Path(args.matter_dir).resolve()
    manifest = load_manifest(matter_dir)
    rows = load_documents(matter_dir)
    output_path = Path(args.output_file)
    if not output_path.exists():
        print(f"ERROR: output file not found: {output_path}")
        return 2
    text = _read_text_best_effort(output_path)
    registered = set(manifest.get("bates_prefixes", []))
    cache = _index_dir(matter_dir) / TEXT_CACHE_DIRNAME

    failures: List[str] = []
    warnings: List[str] = []
    entries = 0

    # Collect dated entries from BOTH layouts:
    #  - "Date: YYYY-MM-DD" blocks (segment runs to the next Date: line)
    #  - markdown table rows "| YYYY-MM-DD | event | TVRR-PROD-000001 |"
    # A table-layout chronology previously verified zero rows and PASSed
    # vacuously.
    matches = list(_CHRONO_DATE_RE.finditer(text))
    dated_segments: List[Tuple[str, str]] = []
    for i, m in enumerate(matches):
        seg_end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        dated_segments.append((m.group(1), text[m.start():seg_end]))
    table_rows = list(_CHRONO_TABLE_ROW_RE.finditer(text))
    for m in table_rows:
        dated_segments.append((m.group(1), m.group(0)))

    for iso, segment in dated_segments:
        cites = [(p, n) for p, n in _extract_citations(segment) if p in registered]
        if not cites:
            # Fail-closed: a dated chronology row without a same-matter Source
            # citation is not "unverifiable" — it is incomplete for handoff.
            if not getattr(args, "allow_uncited", False):
                failures.append(
                    f"{iso} entry: dated chronology row has no same-matter "
                    f"Bates Source citation"
                )
            continue
        entries += 1
        entry_label = f"{iso} entry"
        date_verified = False
        for prefix, number in cites:
            doc = _resolve_bates(rows, prefix, number)
            if doc is None:
                failures.append(f"{entry_label}: unresolved citation "
                                f"{prefix}-{number:06d}")
                continue
            fp = cache / f"{doc['sha256']}.txt"
            if not fp.exists() or doc.get("text_extractable") in ("none", "unsupported"):
                warnings.append(f"{entry_label}: cited doc {doc['relpath']} is not "
                                f"text-verifiable (unreadable/no cache) — manual check")
                continue
            doc_text = fp.read_text(encoding="utf-8")
            if any(v in doc_text for v in _date_variants(iso)):
                date_verified = True
        if cites and not date_verified and not any(
            w.startswith(entry_label) for w in warnings
        ):
            warnings.append(f"{entry_label}: date {iso} not found in any cited "
                            f"document — verify the event date against sources")

    # Vacuous chronology: dated rows present but none checked → FAIL.
    # Also FAIL when a Chronology heading exists but zero parseable dates
    # (reformatting to dodge the parser must not greenwash — red-team P1-7).
    has_chrono_heading = bool(
        re.search(r"(?im)^#{1,3}\s+.*\bchronology\b", text)
    )
    if dated_segments and entries == 0 and not failures and not getattr(args, "allow_uncited", False):
        failures.append(
            "chronology has Date: rows but none were citation-verifiable "
            "(vacuous verify-chronology PASS refused)"
        )
    elif (
        has_chrono_heading
        and not dated_segments
        and entries == 0
        and not failures
        and not getattr(args, "allow_empty_chronology", False)
    ):
        failures.append(
            "Chronology section present but no parseable dated rows "
            "(vacuous verify-chronology PASS refused; use --allow-empty-chronology "
            "only for drafts that intentionally omit dated entries)"
        )

    passed = not failures and (not args.strict or not warnings)
    report = {
        "output_file": str(output_path),
        "matter_id": manifest["matter_id"],
        "entries_checked": entries,
        "failures": failures,
        "warnings": warnings,
        "pass": passed,
    }
    if args.json:
        print(json.dumps(report, indent=2))
    else:
        print(f"verify-chronology: {entries} dated+cited entries checked "
              f"against matter '{manifest['matter_id']}'.")
        for f_ in failures:
            print(f"  FAIL: {f_}")
        for w in warnings:
            print(f"  WARN: {w}")
        print("PASS" if passed else "FAIL")
    return 0 if passed else 1


# ── entity management ────────────────────────────────────────────────────────

def cmd_add_entity(args) -> int:
    matter_dir = Path(args.matter_dir).resolve()
    load_manifest(matter_dir)
    entities = load_entities(matter_dir)
    key = _normalize_identifier(args.name)
    ent = entities.setdefault(
        key, {"display": args.name, "aliases": [], "role": args.role or "",
              "origin": "manual", "sources": {}}
    )
    for alias in args.alias or []:
        if alias not in ent["aliases"]:
            ent["aliases"].append(alias)
    if args.role:
        ent["role"] = args.role
    save_entities(matter_dir, entities)
    print(f"Registered entity '{args.name}' ({len(ent['aliases'])} aliases) "
          f"for matter.")
    return 0


# ── selftest ────────────────────────────────────────────────────────────────

def cmd_selftest(args) -> int:
    import tempfile
    ok = True

    def check(name, cond):
        nonlocal ok
        print(f"  {'PASS' if cond else 'FAIL'}: {name}")
        ok = ok and bool(cond)

    print("casegraph selftest")
    # Bates parsing
    check("bates range parse",
          parse_bates_range("TVRR-PROD-000001 - TVRR-PROD-000004") == ("TVRR-PROD", 1, 4))
    check("bates 'through' parse",
          parse_bates_range("ACME-000010 through ACME-000012") == ("ACME", 10, 12))
    check("bates from filename", bates_from_filename("TVRR-PROD-000123.pdf") == ("TVRR-PROD", 123))
    check("normalize homoglyph/case",
          _normalize_identifier("Ｊ.Ｔ.") == _normalize_identifier("j t"))

    with tempfile.TemporaryDirectory(prefix="casegraph_selftest_") as td:
        matter = Path(td) / "matter"
        docs = matter / "production"
        docs.mkdir(parents=True)
        (docs / "TVRR-PROD-000001.md").write_text(
            "**Bates Range:** TVRR-PROD-000001 - TVRR-PROD-000002\n"
            "**Author:** R.K., Trainmaster, Test Valley Railroad\n"
            "**Date:** 2024-11-13\n\nThe conductor reported an unsafe coupling procedure.\n",
            encoding="utf-8")
        ns = argparse.Namespace(matter_dir=str(matter), matter_id="SELFTEST",
                                bates_prefix=["TVRR-PROD"], force=False)
        check("init", cmd_init(ns) == 0)
        nb = argparse.Namespace(matter_dir=str(matter), no_text_cache=False, json=True)
        import io, contextlib
        buf = io.StringIO()
        with contextlib.redirect_stdout(buf):
            rc = cmd_build(nb)
        check("build", rc == 0)
        rows = load_documents(matter)
        check("bates indexed", rows and rows[0]["bates_prefix"] == "TVRR-PROD")
        good = Path(td) / "out_good.md"
        good.write_text("Fact cited to TVRR-PROD-000002.\n", encoding="utf-8")
        bad = Path(td) / "out_bad.md"
        bad.write_text("Fact cited to TVRR-PROD-000099 and NORF-PROD-000001.\n", encoding="utf-8")
        nv = argparse.Namespace(matter_dir=str(matter), output_file=str(good),
                                quotes=True, no_quotes=False, allow_empty=False,
                                allow_declared_gaps=False, json=True)
        with contextlib.redirect_stdout(buf):
            rc_good = cmd_verify_cites(nv)
            nv.output_file = str(bad)
            rc_bad = cmd_verify_cites(nv)
            empty = Path(td) / "out_empty.md"
            empty.write_text("Narrative with no Bates citations at all.\n", encoding="utf-8")
            nv.output_file = str(empty)
            rc_empty = cmd_verify_cites(nv)
            nv.allow_empty = True
            rc_empty_ok = cmd_verify_cites(nv)
            ni = argparse.Namespace(matter_dir=str(matter), output_file=str(bad),
                                    fingerprints=None, strict=False, json=True)
            rc_iso_bad = cmd_check_isolation(ni)
            ni.output_file = str(good)
            rc_iso_good = cmd_check_isolation(ni)
        check("verify-cites pass", rc_good == 0)
        check("verify-cites fail on out-of-range", rc_bad == 1)
        check("verify-cites fail on empty cites", rc_empty == 1)
        check("verify-cites allow-empty", rc_empty_ok == 0)
        check("isolation fail on foreign prefix", rc_iso_bad == 1)
        check("isolation pass on clean output", rc_iso_good == 0)
        check("allowlist present", _allowlist_path().exists() and bool(_load_allowlist()))
    print("selftest:", "PASS" if ok else "FAIL")
    return 0 if ok else 1


# ── main ────────────────────────────────────────────────────────────────────

def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(prog="casegraph", description=__doc__)
    sub = parser.add_subparsers(dest="cmd", required=True)

    p = sub.add_parser("init", help="initialize a matter index")
    p.add_argument("matter_dir")
    p.add_argument("--matter-id", required=True)
    p.add_argument("--bates-prefix", action="append", default=[])
    p.add_argument("--force", action="store_true")
    p.set_defaults(fn=cmd_init)

    p = sub.add_parser("build", help="incremental scan + index")
    p.add_argument("matter_dir")
    p.add_argument("--no-text-cache", action="store_true")
    p.add_argument("--json", action="store_true")
    p.set_defaults(fn=cmd_build)

    p = sub.add_parser("status", help="staleness check (exit 1 if stale)")
    p.add_argument("matter_dir")
    p.add_argument("--deep", action="store_true", help="hash-verify unchanged-looking files")
    p.add_argument("--json", action="store_true")
    p.set_defaults(fn=cmd_status)

    p = sub.add_parser("query", help="index lookups")
    p.add_argument("matter_dir")
    p.add_argument("--bates")
    p.add_argument("--doc")
    p.add_argument("--entity")
    p.add_argument("--grep")
    p.set_defaults(fn=cmd_query)

    p = sub.add_parser("verify-cites", help="citations in output must resolve (exit 1 on failure)")
    p.add_argument("matter_dir")
    p.add_argument("output_file")
    p.add_argument("--quotes", action="store_true", default=True,
                   help="verify quoted strings appear in corpus (default: on)")
    p.add_argument("--no-quotes", action="store_true",
                   help="skip quote verification (draft passes only)")
    p.add_argument("--allow-empty", action="store_true",
                   help="allow zero same-matter citations (default: fail closed)")
    p.add_argument("--allow-declared-gaps", action="store_true",
                   help="treat cover-letter-declared but unindexed Bates as INFO "
                        "(default: FAIL — prevents cite laundering)")
    p.add_argument("--json", action="store_true")
    p.set_defaults(fn=cmd_verify_cites)

    p = sub.add_parser("verify-chronology",
                       help="dated events must trace to docs containing the date (exit 1 on FAIL)")
    p.add_argument("matter_dir")
    p.add_argument("output_file")
    p.add_argument("--strict", action="store_true", help="unresolved WARNs also fail")
    p.add_argument("--allow-uncited", action="store_true",
                   help="skip Date: rows that lack same-matter Bates Sources")
    p.add_argument("--allow-empty-chronology", action="store_true",
                   help="allow Chronology heading with no parseable dated rows")
    p.add_argument("--json", action="store_true")
    p.set_defaults(fn=cmd_verify_chronology)

    p = sub.add_parser("check-isolation", help="cross-matter contamination gate (exit 1 on FAIL)")
    p.add_argument("matter_dir")
    p.add_argument("output_file")
    p.add_argument("--fingerprints", help="shared salted-hash fingerprint store")
    p.add_argument("--strict", action="store_true", help="unresolved WARNs also fail")
    p.add_argument("--json", action="store_true")
    p.set_defaults(fn=cmd_check_isolation)

    p = sub.add_parser("add-entity", help="register an entity for the matter")
    p.add_argument("matter_dir")
    p.add_argument("--name", required=True)
    p.add_argument("--alias", action="append", default=[])
    p.add_argument("--role")
    p.set_defaults(fn=cmd_add_entity)

    p = sub.add_parser("export-fingerprint", help="publish salted identifier hashes")
    p.add_argument("matter_dir")
    p.add_argument("--store", required=True)
    p.set_defaults(fn=cmd_export_fingerprint)

    p = sub.add_parser(
        "export-ocr-queue",
        help="list docs needing OCR/text; write .casegraph/needs_ocr.json (exit 1 if any)",
    )
    p.add_argument("matter_dir")
    p.add_argument("--json", action="store_true")
    p.set_defaults(fn=cmd_export_ocr_queue)

    p = sub.add_parser("selftest", help="offline self-test")
    p.set_defaults(fn=cmd_selftest)

    args = parser.parse_args(argv)
    return args.fn(args)


if __name__ == "__main__":
    sys.exit(main())
